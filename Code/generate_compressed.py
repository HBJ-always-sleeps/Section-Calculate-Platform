from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_run_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_h(doc, text, fn='黑体', sz=16, bd=True, sa=12, sb=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, fn, sz, bd)
    return p

def add_p(doc, text, fi=0.74, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if fi:
        p.paragraph_format.first_line_indent = Cm(fi)
    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', 14, False)
    return p

def add_bp(doc, lead, content, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(lead)
    set_run_font(r1, '仿宋_GB2312', 14, True)
    r2 = p.add_run(content)
    set_run_font(r2, '仿宋_GB2312', 14, False)
    return p

doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3.7)
sec.bottom_margin = Cm(3.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.6)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run('附件2')
set_run_font(run, '黑体', 22, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('"导师带徒"优秀案例')
set_run_font(run, '黑体', 22, True)

add_h(doc, '案例标题', sa=6, sb=12)
add_p(doc, '薪火相传促成长，测量实训筑传承——广西北海项目"导师带徒"优秀案例', fi=0, sa=12)

add_bp(doc, '报送单位：', '中交华南勘察测绘科技有限公司广西北海项目部', sa=18)

# 一、师徒基本情况
add_h(doc, '一、师徒基本情况', sa=10, sb=12)
add_bp(doc, '导师：', '陈杰，广西北海项目测量部，测量部长。长期从事海洋测绘与工程测量工作，在多波束测深系统、GNSS精密定位、水文测量等领域具有丰富的项目实践经验。', sa=8)
add_bp(doc, '徒弟：', '黄秉俊，广西北海项目测量部，见习生。毕业于中山大学遥感科学与技术专业，2025年7月入职并参加师带徒结对培养。', sa=8)
add_bp(doc, '结对时间：', '2025年7月至2026年7月，协议期限2025年7月13日至2026年7月12日。', sa=8)
add_bp(doc, '培养方向/目标：', '围绕海洋测绘与工程测量技术方向，按照"基础适应—技能深化—技术拓展—综合应用"四阶段递进培养。通过理论学习、现场实操、项目历练相结合的方式，帮助徒弟系统掌握多波束与单波束测深、RTK精密控制测量、Caris水文数据处理、CAD工程算量等核心专业技能，实现从毕业生向合格工程技术人员的角色转变。', sa=8)
add_bp(doc, '案例关键词：', '快速融入、技术攻坚、以干代训、角色转换、亦师亦友。', sa=12)

# 二、主要做法与过程
add_h(doc, '二、主要做法与过程（重点撰写）', sa=10, sb=12)
add_h(doc, '精准制定培养计划：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '师傅充分结合徒弟的专业背景与岗位实际，量身定制了为期一年的个性化培养方案。整个培养路径划分为四个递进阶段：基础适应阶段以熟悉项目环境、牢记安全规范、夯实测量理论基础为主；技能深化阶段重点开展单波束与多波束测深系统操作、RTK控制测量等核心技能训练；技术拓展阶段进一步学习星基RTK应用、多波束仪器配置安装及Caris软件进阶数据处理；综合应用阶段则在师傅指导下参与设备外业调试、内业数据处理与工程算量等实战任务。各阶段均配套明确的学习任务清单、实操考核标准和时间节点要求。', sa=8)

add_h(doc, '创新带教方法：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '师傅始终坚持"讲解—示范—实操—复盘"的闭环教学模式。理论学习通过专题讲解、设备手册研读、规范学习及内部培训相结合的方式开展；现场操作环节亲自示范仪器架设、参数配置、数据采集等关键步骤，反复强调安全底线与操作规范。尤为可贵的是，师傅注重培养徒弟的独立思考能力，要求徒弟在遇到技术问题时先独立排查、查阅资料，再进行针对性指导。此外，师徒之间建立了常态化的谈心谈话机制，自2025年7月至2026年3月累计开展谈心谈话9次，内容涵盖技术难点剖析、职业规划指导、心理状态调适等方面。', sa=8)

add_h(doc, '徒弟学习与实践：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '在整个培养周期内，徒弟始终保持积极主动的学习态度和刻苦钻研的进取精神。外业方面，先后参与了测区多波束数据采集、吹填区围堰RTK加密控制测量等任务，熟练掌握了项目指定思拓力RTK设备的完整作业流程。内业方面，在师傅指导下完成了Caris数据假水删除处理、三维水深图制作、多波束测量记录表整理、断面数据砂类算量统计等工作。针对Caris软件操作尚未完全熟练的情况，徒弟主动录制操作视频反复练习；面对从遥感宏观视角向测绘微观落实的思维转变，通过大量现场实践逐步完成了认知转换。', sa=8)

add_h(doc, '过程管理督导：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '项目部建立了"现场记录—内业核对—成果输出"的全流程质量管控机制。师傅定期检查徒弟的学习笔记、外业记录表和内业处理成果，对发现的问题当场指出、及时纠正。在CAD数据处理与算量阶段，徒弟曾因坐标单位未统一导致面积计算出现偏差，经师傅指导后立即建立交叉验证机制，并系统整理形成《常见错误清单》和《算量流程说明》。培养过程实行季度考核制度，考核维度涵盖学习任务完成情况、知识掌握程度、工作主动性、问题发现与改进能力等十个方面，确保培养质量可量化、可追溯。', sa=12)

# 三、成果与成效
add_h(doc, '三、成果与成效（重点撰写，尽量量化）', sa=10, sb=12)
add_h(doc, '徒弟技能提升方面：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '经过系统培养，徒弟已初步掌握多波束测深系统的基本操作与现场数据采集、RTK精密控制测量全流程（含星基RTK与地基增强两种模式的适用场景判断与参数配置）、Caris水文数据处理软件的基础操作与假水数据剔除、CAD工程制图与断面提取及算量、测量数据记录整理与图表制作等核心技能。具体可视化成果包括Caris数据假水删除处理成果、测区三维水深图、多波束测量记录表、吹填区围堰RTK加密测量成果表、断面数据砂类算量统计表等。徒弟已能够独立完成多个控制点的RTK测量与记录表填写，并能在复杂环境下自主完成仪器初始化与精度监控。', sa=8)

add_h(doc, '徒弟业绩表现方面：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '在角色定位上，徒弟实现了从"被动跟学"到"主动参与"再到"部分独立完成"的阶梯式跨越。第一季度以夯实理论基础和掌握RTK基本操作为主；第二季度进入技术拓展与综合应用阶段，使用CAD软件完成外业成果整理、图层规范、要素检查、断面提取与面积统计，并通过两种方法进行交叉验证，成功发现并修正了因单位不统一导致的面积计算错误，避免了工程计量偏差。徒弟还将星基RTK的精度范围、适用工况以及与地基增强的切换原则系统整理为《星基RTK作业指导卡》，为后续同类作业提供了标准化参考。', sa=8)

add_h(doc, '徒弟综合素质方面：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '责任心显著增强，深刻体会到外业采集的微小疏漏会给内业数据处理带来成倍返工，逐步养成了"一次做对、全程留痕"的严谨作风。团队协作意识有效提升，能够主动适应项目部"分工明确、随时补位"的工作节奏。解决问题能力明显提高，面对围堰转角遮挡导致卫星信号频繁失锁的实际困难，在师傅指导下学会了从星基增强模式切换至地基增强模式，理解了两种改正数传播机制的差异及适用边界，并总结出"开阔区初始化—遮挡区复测—关键点位往返检核"的实操经验。', sa=8)

add_h(doc, '徒弟业务促进方面：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '徒弟参与完成的吹填区围堰加密控制测量成果、多波束数据采集与处理成果，为项目部相关工程的施工放样、工程量计算和质量验收提供了准确的基础数据支撑。通过建立"现场记录—内业核对—成果输出"的标准化工作流程，徒弟将常用坐标参数、检查项清单和常见错误排查要点整理成册，形成了可在项目内部推广应用的作业模板。在Caris数据处理中，徒弟针对边缘波束信噪比低、假水点识别困难等问题，与师傅共同探讨出"优先检查大角度边缘波束—结合等深线趋势与发射接收角度综合判断—可疑区域单独标注复测"的三步处理法，有效提升了数据处理效率和成果可靠性。', sa=12)

# 四、经验启示与展望
add_h(doc, '四、经验启示与展望', sa=10, sb=12)
add_h(doc, '成功关键因素分析：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '本案例取得阶段性成效，得益于三个方面的有机结合。一是师徒磨合顺畅，师傅能够准确把握徒弟遥感专业背景与测绘岗位需求的契合点，因材施教地将遥感影像分析中的空间思维与测绘点位精度控制相结合，帮助徒弟顺利完成知识体系的跨界迁移。二是培养计划具有清晰的阶段性和可操作性，从基础适应到综合应用的四个阶段层层递进，每阶段均设定了明确的能力目标和验收标准。三是项目部实践氛围浓厚，北海项目部虽然节奏快、事务杂，但团队协作紧密、容错空间充足，为徒弟提供了大量真实的项目练兵机会。', sa=8)

add_h(doc, '总结归纳：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '一是闭环教学模式值得推广。"讲解—示范—实操—复盘"的四步法将知识传授、技能训练与反思改进融为一体，尤其适用于测绘这类对操作精度和规范意识要求极高的技术岗位，能够有效缩短新员工的技能爬坡期。二是问题导向的培养方式成效显著。鼓励徒弟先独立排查问题、查阅资料，师傅再进行针对性点拨，既保护了徒弟的探索积极性，又避免了低效重复。三是定期谈心机制不可忽视。将技术指导与职业规划、心理疏导有机融合，有助于徒弟完成从"校园学习者"到"岗位责任者"的心态转变。', sa=8)

add_h(doc, '相关展望：', '楷体_GB2312', 15, True, sa=6, sb=8)
add_p(doc, '基于本案例的实践探索，对公司导师带徒工作提出以下优化建议。第一，建议在条件允许的范围内增加常见仪器故障模拟与数据异常排查的专项训练，帮助新员工在低风险环境中积累排障经验。第二，建议将项目实践中的典型技术难点和解决经验整理形成"微案例库"，按照设备操作、数据处理、质量控制等维度分类归档，便于后续批次徒弟快速检索学习、复盘借鉴。第三，建议适时组织跨项目的技术交流活动，让处于不同培养阶段的徒弟相互分享成长心得，让经验丰富的师傅之间切磋带教方法，进一步拓宽青年员工的技术视野和成长通道。', sa=12)

# 五、导师寄语/徒弟感言
add_h(doc, '五、导师寄语/徒弟感言(选填，可填简短一句话或一段话)', sa=10, sb=12)
add_p(doc, '（师徒双方可根据实际情况填写，此处暂空待补。）', sa=12)

# 备注
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('（备注：正文字数控制在2500字以内，符合模板建议范围。）')
set_run_font(run, '仿宋_GB2312', 12, False)

output_path = r'C:\Users\训教\Desktop\师带徒\附件2_导师带徒优秀案例_黄秉俊_压缩版.docx'
doc.save(output_path)
print('Saved to:', output_path)
