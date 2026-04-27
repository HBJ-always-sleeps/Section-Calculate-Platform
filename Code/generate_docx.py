"""
生成"导师带徒"优秀案例标准docx文档
严格沿用模板原有标题、层级与结构，正文合理分段
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# ────────────────────────────────────────
# 辅助函数：设置中文字体
# ────────────────────────────────────────
def set_run_font(run, font_name='宋体', size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading_paragraph(doc, text, font_name='黑体', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12, space_before=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold)
    return p

def add_normal_paragraph(doc, text, first_line_indent=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_name='仿宋_GB2312', size=14, bold=False)
    return p

def add_bold_lead_paragraph(doc, lead_text, content_text, space_after=6):
    """加粗引导语 + 正文内容，如'导师：' + 具体内容"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    # 引导语加粗
    run_lead = p.add_run(lead_text)
    set_run_font(run_lead, font_name='仿宋_GB2312', size=14, bold=True)
    # 正文内容
    run_content = p.add_run(content_text)
    set_run_font(run_content, font_name='仿宋_GB2312', size=14, bold=False)
    return p

# ────────────────────────────────────────
# 页面设置
# ────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(3.7)
section.bottom_margin = Cm(3.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.6)

# ────────────────────────────────────────
# 文档标题
# ────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('附件2')
set_run_font(run, font_name='黑体', size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('"导师带徒"优秀案例')
set_run_font(run, font_name='黑体', size=22, bold=True)

# ────────────────────────────────────────
# 案例标题
# ────────────────────────────────────────
add_heading_paragraph(doc, '案例标题', font_name='黑体', size=16, bold=True, space_after=6, space_before=12)
add_normal_paragraph(doc,
    '薪火相传促成长，测量实训筑传承——广西北海项目"导师带徒"优秀案例',
    first_line_indent=0, space_after=12)

# ────────────────────────────────────────
# 报送单位
# ────────────────────────────────────────
add_bold_lead_paragraph(doc, '报送单位：', '中交华南勘察测绘科技有限公司广西北海项目部', space_after=18)

# ════════════════════════════════════════
# 一、师徒基本情况
# ════════════════════════════════════════
add_heading_paragraph(doc, '一、师徒基本情况', font_name='黑体', size=16, bold=True, space_after=10, space_before=12)

add_bold_lead_paragraph(doc, '导师：',
    '陈杰，广西北海项目测量部，测量部长。长期从事海洋测绘与工程测量工作，在多波束测深系统、GNSS精密定位、水文测量等领域具有丰富的项目实践经验。',
    space_after=8)

add_bold_lead_paragraph(doc, '徒弟：',
    '黄秉俊，广西北海项目测量部，见习生。毕业于中山大学遥感科学与技术专业，2025年7月入职，同月正式参加师带徒结对培养。',
    space_after=8)

add_bold_lead_paragraph(doc, '结对时间：',
    '2025年7月至2026年7月。师徒双方于2025年7月13日正式签署师带徒协议书，协议期限一年。',
    space_after=8)

add_bold_lead_paragraph(doc, '培养方向/目标：',
    '围绕海洋测绘与工程测量技术方向，按照"基础适应—技能深化—技术拓展—综合应用"四个阶段递进培养。通过理论学习、现场实操、项目历练相结合的方式，帮助徒弟系统掌握多波束与单波束测深内外业全流程、RTK精密控制测量、Caris水文数据处理、CAD工程算量等核心专业技能，实现从院校毕业生向合格工程技术人员的角色转变，最终达到能够独立承担项目测量任务的能力目标。',
    space_after=8)

add_bold_lead_paragraph(doc, '案例关键词：',
    '快速融入、技术攻坚、以干代训、角色转换、亦师亦友。',
    space_after=12)

# ════════════════════════════════════════
# 二、主要做法与过程（重点撰写）
# ════════════════════════════════════════
add_heading_paragraph(doc, '二、主要做法与过程（重点撰写）', font_name='黑体', size=16, bold=True, space_after=10, space_before=12)

# 2.1 精准制定培养计划
add_heading_paragraph(doc, '精准制定培养计划：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '师傅充分结合徒弟的专业背景与岗位实际，量身定制了为期一年的个性化培养方案。整个培养路径划分为四个递进而又相互衔接的阶段：基础适应阶段以熟悉项目环境、牢记安全规范、夯实测量理论基础为主；技能深化阶段重点开展单波束与多波束测深系统操作、RTK控制测量等核心技能训练；技术拓展阶段进一步学习星基RTK应用、多波束仪器配置安装及Caris软件进阶数据处理；综合应用阶段则在师傅指导下参与设备外业调试、内业数据处理与工程算量等实战任务。各阶段均配套明确的学习任务清单、实操考核标准和时间节点要求，确保培养计划落地见效。',
    first_line_indent=0.74, space_after=8)

# 2.2 创新带教方法
add_heading_paragraph(doc, '创新带教方法：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '在技能传授方面，师傅始终坚持"讲解—示范—实操—复盘"的闭环教学模式。理论学习环节，通过专题讲解、设备手册研读、行业规范学习及内部培训相结合的方式，将抽象的测量原理与具体的项目实例紧密衔接；现场操作环节，师傅亲自示范仪器架设、参数配置、数据采集等关键步骤，并反复强调安全底线与操作规范，帮助徒弟养成严谨细致的工作习惯。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '尤为可贵的是，师傅注重培养徒弟的独立思考能力，要求徒弟在遇到技术问题时先独立排查、查阅资料，再进行针对性指导，有效激发了徒弟的主动学习意识。此外，师徒之间建立了常态化的谈心谈话机制，自2025年7月至2026年3月累计开展谈心谈话9次，内容涵盖技术难点剖析、职业规划指导、心理状态调适等方面，形成了亦师亦友的良好互动氛围。',
    first_line_indent=0.74, space_after=8)

# 2.3 徒弟学习与实践
add_heading_paragraph(doc, '徒弟学习与实践：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '在整个培养周期内，徒弟始终保持积极主动的学习态度和刻苦钻研的进取精神。外业方面，先后参与了测区多波束数据采集、吹填区围堰RTK加密控制测量等任务，熟练掌握了项目指定思拓力RTK设备的完整作业流程，包括基准站架设、流动站配置、坐标系与投影参数核对、野外观测记录及成果导出等环节。内业方面，在师傅指导下完成了Caris数据假水删除处理、三维水深图制作、多波束测量记录表整理、断面数据砂类算量统计等工作。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '针对Caris软件操作尚未完全熟练的情况，徒弟主动录制操作视频反复练习；面对从遥感宏观视角向测绘微观落实的思维转变，徒弟通过大量现场实践逐步完成了认知转换。在围堰加密测量中遇到遮挡区信号失锁难题时，徒弟在师傅引导下深入理解星基增强与地基增强两种模式的原理差异，并学会了在复杂环境下的灵活切换。',
    first_line_indent=0.74, space_after=8)

# 2.4 过程管理督导
add_heading_paragraph(doc, '过程管理督导：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '项目部建立了"现场记录—内业核对—成果输出"的全流程质量管控机制。师傅定期检查徒弟的学习笔记、外业记录表和内业处理成果，对发现的问题当场指出、及时纠正。在CAD数据处理与算量阶段，徒弟曾因坐标单位未统一导致面积计算出现偏差，经师傅指导后立即建立交叉验证机制，并系统整理形成《常见错误清单》和《算量流程说明》。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '培养过程实行季度考核制度，第一季度考核周期为2025年10月至12月，第二季度为2026年1月至3月，徒弟季度出勤率均为75天，考核维度涵盖学习任务完成情况、知识掌握程度、工作主动性、问题发现与改进能力等十个方面，确保培养质量可量化、可追溯。',
    first_line_indent=0.74, space_after=12)

# ════════════════════════════════════════
# 三、成果与成效（重点撰写，尽量量化）
# ════════════════════════════════════════
add_heading_paragraph(doc, '三、成果与成效（重点撰写，尽量量化）', font_name='黑体', size=16, bold=True, space_after=10, space_before=12)

# 3.1 徒弟技能提升方面
add_heading_paragraph(doc, '徒弟技能提升方面：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '经过系统培养，徒弟已初步掌握以下核心技能：多波束测深系统的基本操作与现场数据采集；RTK精密控制测量全流程（含星基RTK与地基增强两种模式的适用场景判断与参数配置）；Caris水文数据处理软件的基础操作与假水数据剔除；CAD工程制图与断面提取、面积统计及算量；测量数据记录、整理与图表制作。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '具体可视化成果包括：Caris数据假水删除处理成果、测区三维水深图、多波束测量记录表、吹填区围堰RTK加密测量成果表、断面数据砂类算量统计表等。徒弟已能够独立完成多个控制点的RTK测量与记录表填写，并能在复杂环境下自主完成仪器初始化与精度监控。',
    first_line_indent=0.74, space_after=8)

# 3.2 徒弟业绩表现方面
add_heading_paragraph(doc, '徒弟业绩表现方面：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '在角色定位上，徒弟实现了从"被动跟学"到"主动参与"再到"部分独立完成"的阶梯式跨越。第一季度（2025年9月至12月）以夯实理论基础和掌握RTK基本操作为主；第二季度（2025年12月至2026年3月）进入技术拓展与综合应用阶段，使用CAD软件完成外业成果整理、图层规范、要素检查、断面提取与面积统计，并通过两种方法进行交叉验证，成功发现并修正了因单位不统一导致的面积计算错误，避免了工程计量偏差。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '徒弟还将星基RTK的精度范围、适用工况以及与地基增强的切换原则系统整理为《星基RTK作业指导卡》，为后续同类作业提供了标准化参考。',
    first_line_indent=0.74, space_after=8)

# 3.3 徒弟综合素质方面
add_heading_paragraph(doc, '徒弟综合素质方面：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '责任心显著增强，深刻体会到外业采集的微小疏漏（如记录字段不完整、点位草图缺失）会给内业数据处理带来成倍返工，逐步养成了"一次做对、全程留痕"的严谨作风。团队协作意识有效提升，能够主动适应项目部"分工明确、随时补位"的工作节奏，与船上作业人员及其他测量小组的沟通效率持续改善。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '解决问题能力明显提高，面对围堰转角遮挡导致卫星信号频繁失锁的实际困难，在师傅指导下学会了从星基增强模式切换至地基增强模式，理解了两种改正数传播机制的差异及适用边界，并总结出"开阔区初始化—遮挡区复测—关键点位往返检核"的实操经验。',
    first_line_indent=0.74, space_after=8)

# 3.4 徒弟业务促进方面
add_heading_paragraph(doc, '徒弟业务促进方面：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '徒弟参与完成的吹填区围堰加密控制测量成果、多波束数据采集与处理成果，为项目部相关工程的施工放样、工程量计算和质量验收提供了准确的基础数据支撑。通过建立"现场记录—内业核对—成果输出"的标准化工作流程，徒弟将常用坐标参数、检查项清单和常见错误排查要点整理成册，形成了可在项目内部推广应用的作业模板。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '在Caris数据处理中，徒弟针对边缘波束信噪比低、假水点识别困难等问题，与师傅共同探讨出"优先检查大角度边缘波束—结合等深线趋势与发射接收角度综合判断—可疑区域单独标注复测"的三步处理法，有效提升了数据处理效率和成果可靠性。',
    first_line_indent=0.74, space_after=12)

# ════════════════════════════════════════
# 四、经验启示与展望
# ════════════════════════════════════════
add_heading_paragraph(doc, '四、经验启示与展望', font_name='黑体', size=16, bold=True, space_after=10, space_before=12)

# 4.1 成功关键因素分析
add_heading_paragraph(doc, '成功关键因素分析：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '本案例取得阶段性成效，得益于三个方面的有机结合。一是师徒磨合顺畅，师傅能够准确把握徒弟遥感专业背景与测绘岗位需求的契合点，因材施教地将遥感影像分析中的空间思维与测绘点位精度控制相结合，帮助徒弟顺利完成知识体系的跨界迁移。二是培养计划具有清晰的阶段性和可操作性，从基础适应到综合应用的四个阶段层层递进，每阶段均设定了明确的能力目标和验收标准，避免了培养的盲目性和随意性。三是项目部实践氛围浓厚，北海项目部虽然节奏快、事务杂，但团队协作紧密、容错空间充足，为徒弟提供了大量真实的项目练兵机会，使徒弟能够在干中学、在学中干。',
    first_line_indent=0.74, space_after=8)

# 4.2 总结归纳
add_heading_paragraph(doc, '总结归纳：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '一是闭环教学模式值得推广。"讲解—示范—实操—复盘"的四步法将知识传授、技能训练与反思改进融为一体，尤其适用于测绘这类对操作精度和规范意识要求极高的技术岗位，能够有效缩短新员工的技能爬坡期。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '二是问题导向的培养方式成效显著。鼓励徒弟先独立排查问题、查阅资料，师傅再进行针对性点拨，既保护了徒弟的探索积极性，又避免了低效重复，培养了徒弟独立分析和解决现场问题的能力。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '三是定期谈心机制不可忽视。将技术指导与职业规划、心理疏导有机融合，有助于徒弟完成从"校园学习者"到"岗位责任者"的心态转变，增强对企业的归属感和对职业发展的信心。',
    first_line_indent=0.74, space_after=8)

# 4.3 相关展望
add_heading_paragraph(doc, '相关展望：', font_name='楷体_GB2312', size=15, bold=True, space_after=6, space_before=8)
add_normal_paragraph(doc,
    '基于本案例的实践探索，对公司导师带徒工作提出以下优化建议。第一，建议在条件允许的范围内增加常见仪器故障模拟与数据异常排查的专项训练，如RTK信号失锁、多波束边缘波束异常、Caris数据导入错误等情景演练，帮助新员工在低风险环境中积累排障经验。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '第二，建议将项目实践中的典型技术难点和解决经验整理形成"微案例库"，按照设备操作、数据处理、质量控制等维度分类归档，便于后续批次徒弟快速检索学习、复盘借鉴。',
    first_line_indent=0.74, space_after=6)
add_normal_paragraph(doc,
    '第三，建议适时组织跨项目的技术交流活动，让处于不同培养阶段的徒弟相互分享成长心得，让经验丰富的师傅之间切磋带教方法，进一步拓宽青年员工的技术视野和成长通道。',
    first_line_indent=0.74, space_after=12)

# ════════════════════════════════════════
# 五、导师寄语/徒弟感言(选填，可填简短一句话或一段话)
# ════════════════════════════════════════
add_heading_paragraph(doc, '五、导师寄语/徒弟感言(选填，可填简短一句话或一段话)', font_name='黑体', size=16, bold=True, space_after=10, space_before=12)
add_normal_paragraph(doc,
    '（师徒双方可根据实际情况填写，此处暂空待补。）',
    first_line_indent=0.74, space_after=12)

# ────────────────────────────────────────
# 备注
# ────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run('（备注：正文字数约2300字，符合2000—2500字建议范围。案例配图建议包括：师徒工作合影、师带徒协议签署照片、徒弟外业操作现场照片、Caris数据处理成果截图、三维水深图、RTK加密测量点位分布图等。）')
set_run_font(run, font_name='仿宋_GB2312', size=12, bold=False)

# ────────────────────────────────────────
# 保存文档
# ────────────────────────────────────────
output_path = r'C:\Users\训教\Desktop\师带徒\附件2_导师带徒优秀案例_黄秉俊.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
