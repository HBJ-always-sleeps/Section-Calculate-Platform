import docx
import os
import sys

# Fix encoding
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

dir_path = r"C:\Users\训教\Desktop\师带徒"
files = os.listdir(dir_path)

# Find the template file - exclude ~$ temp files
template_file = None
for f in files:
    if '模板' in f and f.endswith('.docx') and not f.startswith('~'):
        template_file = f
        break

output_lines = []

if template_file:
    template_path = os.path.join(dir_path, template_file)
    doc = docx.Document(template_path)
    output_lines.append("=== 模板段落内容 ===")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.replace('\u200b', '').strip()
        if text:
            output_lines.append(f"[{i}] Style: {p.style.name} | Text: {text}")
    output_lines.append("\n=== 表格内容 ===")
    for ti, table in enumerate(doc.tables):
        output_lines.append(f"\n--- Table {ti} ---")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                ct = cell.text.replace('\u200b', '').strip()
                if ct:
                    output_lines.append(f"  Row {ri}, Col {ci}: {ct}")

# Read the notice file
notice_file = None
for f in files:
    if '通知' in f and not f.startswith('~'):
        notice_file = f
        break

if notice_file:
    notice_path = os.path.join(dir_path, notice_file)
    output_lines.append("\n\n=== 通知文件内容 ===")
    try:
        doc2 = docx.Document(notice_path)
        for i, p in enumerate(doc2.paragraphs):
            text = p.text.replace('\u200b', '').strip()
            if text:
                output_lines.append(f"[{i}] {text}")
        for ti, table in enumerate(doc2.tables):
            output_lines.append(f"\n--- Table {ti} ---")
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    ct = cell.text.replace('\u200b', '').strip()
                    if ct:
                        output_lines.append(f"  Row {ri}, Col {ci}: {ct}")
    except Exception as e:
        output_lines.append(f"Cannot read as docx: {e}")

# Read the 谈话 docx
tanhua_file = None
for f in files:
    if '谈话' in f and f.endswith('.docx') and 'backup' not in f and not f.startswith('~'):
        tanhua_file = f
        break

if tanhua_file:
    tanhua_path = os.path.join(dir_path, tanhua_file)
    output_lines.append("\n\n=== 谈话文件内容 ===")
    try:
        doc4 = docx.Document(tanhua_path)
        for i, p in enumerate(doc4.paragraphs):
            text = p.text.replace('\u200b', '').strip()
            if text:
                output_lines.append(f"[{i}] {text}")
        for ti, table in enumerate(doc4.tables):
            output_lines.append(f"\n--- Table {ti} ---")
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    ct = cell.text.replace('\u200b', '').strip()
                    if ct:
                        output_lines.append(f"  Row {ri}, Col {ci}: {ct}")
    except Exception as e:
        output_lines.append(f"Cannot read: {e}")

# Read the 协议书 file
xieyi_file = None
for f in files:
    if '协议书' in f and not f.startswith('~'):
        xieyi_file = f
        break

if xieyi_file:
    xieyi_path = os.path.join(dir_path, xieyi_file)
    output_lines.append("\n\n=== 协议书文件内容 ===")
    try:
        doc5 = docx.Document(xieyi_path)
        for i, p in enumerate(doc5.paragraphs):
            text = p.text.replace('\u200b', '').strip()
            if text:
                output_lines.append(f"[{i}] {text}")
        for ti, table in enumerate(doc5.tables):
            output_lines.append(f"\n--- Table {ti} ---")
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    ct = cell.text.replace('\u200b', '').strip()
                    if ct:
                        output_lines.append(f"  Row {ri}, Col {ci}: {ct}")
    except Exception as e:
        output_lines.append(f"Cannot read as docx: {e}")

# Write output to file
output_path = os.path.join(dir_path, 'extracted_content.txt')
with open(output_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output_lines))

print(f"Output written to: {output_path}")
print(f"Total lines: {len(output_lines)}")
